import docx  # wymaga instalacji modułu python-docx
# https://python-docx.readthedocs.io/en/latest/index.html

pbp_fname = 'dane/pbp.docx'
doc = docx.Document(pbp_fname)

print(len(doc.paragraphs))
print(len(doc.tables))
second_table = doc.tables[1]
title_column = second_table.columns[0]
for cell in title_column.cells:
    print(cell.text)
